from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader

from txt_key_generator.schemas import (
    FileMetadata,
    FileReadResult,
    GenerationKeys,
    InputKind,
    KeyValidationResult,
    RouteResult,
)


SUPPORTED_EXTENSIONS = {"txt", "md", "docx", "pdf"}


class FileReaderSkill:
    """Read supported Telegram document payloads into plain text."""

    supported_extensions = SUPPORTED_EXTENSIONS

    def read_bytes(
        self,
        content: bytes,
        filename: str,
        mime_type: str | None = None,
    ) -> FileReadResult:
        extension = Path(filename).suffix.lower().lstrip(".")
        metadata = FileMetadata(
            filename=filename,
            extension=extension,
            mime_type=mime_type,
            size_bytes=len(content),
        )

        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file extension: {extension or 'unknown'}")

        if extension in {"txt", "md"}:
            text = self._read_text(content)
        elif extension == "docx":
            text = self._read_docx(content)
        else:
            text = self._read_pdf(content)

        text = text.strip()
        if not text:
            raise ValueError("File does not contain extractable text.")

        return FileReadResult(text=text, metadata=metadata)

    def _read_text(self, content: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")

    def _read_docx(self, content: bytes) -> str:
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells: list[str] = []
        for table in document.tables:
            for row in table.rows:
                table_cells.extend(cell.text for cell in row.cells)
        return "\n".join([*paragraphs, *table_cells])

    def _read_pdf(self, content: bytes) -> str:
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)


class InputRouterSkill:
    """Classify incoming Telegram content before orchestration."""

    def route(
        self,
        text: str | None = None,
        filename: str | None = None,
        is_command: bool = False,
        previous_questions_pending: bool = False,
    ) -> RouteResult:
        if is_command:
            return RouteResult(kind=InputKind.COMMAND, reason="Telegram command.")

        if previous_questions_pending and text and text.strip():
            return RouteResult(kind=InputKind.CLARIFICATION, reason="Clarification reply.")

        if filename:
            extension = Path(filename).suffix.lower().lstrip(".")
            if extension in SUPPORTED_EXTENSIONS:
                return RouteResult(
                    kind=InputKind.FILE_REQUEST,
                    reason="Supported document upload.",
                    extension=extension,
                )
            return RouteResult(
                kind=InputKind.UNSUPPORTED_FILE,
                reason="Unsupported document extension.",
                extension=extension or None,
            )

        if text and text.strip():
            return RouteResult(kind=InputKind.TEXT_REQUEST, reason="Text request.")

        return RouteResult(kind=InputKind.EMPTY, reason="No text or document.")


class KeyValidationSkill:
    required_fields = ("topic", "language", "target_word_count")

    def validate(self, keys: GenerationKeys) -> KeyValidationResult:
        missing = [
            field_name
            for field_name in self.required_fields
            if getattr(keys, field_name) in (None, "", [])
        ]
        conflicts = self._find_conflicts(keys)
        questions = self._build_questions(missing, conflicts, keys)
        return KeyValidationResult(
            is_complete=not missing and not conflicts,
            missing_fields=missing,
            conflicts=conflicts,
            clarifying_questions=questions,
        )

    def _find_conflicts(self, keys: GenerationKeys) -> list[str]:
        conflicts: list[str] = []
        include = {value.casefold() for value in keys.must_include}
        avoid = {value.casefold() for value in keys.must_avoid}
        keyword_set = {value.casefold() for value in keys.keywords}
        blocked_keywords = keyword_set & avoid
        include_and_avoid = include & avoid
        if blocked_keywords:
            conflicts.append(
                "Some keywords are also listed in must_avoid: "
                + ", ".join(sorted(blocked_keywords))
            )
        if include_and_avoid:
            conflicts.append(
                "Some phrases are both required and forbidden: "
                + ", ".join(sorted(include_and_avoid))
            )
        return conflicts

    def _build_questions(
        self,
        missing: Iterable[str],
        conflicts: Iterable[str],
        keys: GenerationKeys,
    ) -> list[str]:
        labels = {
            "topic": "Какая тема текста?",
            "language": "На каком языке нужно написать текст?",
            "target_word_count": "Сколько слов должно быть в тексте?",
        }
        questions = [labels[field_name] for field_name in missing]
        if not keys.keywords:
            questions.append("Какие ключевые слова нужно использовать?")
        questions.extend(f"Уточните конфликт: {conflict}" for conflict in conflicts)
        return questions[:5]


class MessageSplitterSkill:
    def __init__(self, max_length: int = 3900) -> None:
        self.max_length = max_length

    def split(self, message: str) -> list[str]:
        if len(message) <= self.max_length:
            return [message]

        chunks: list[str] = []
        current = ""
        for paragraph in message.split("\n"):
            candidate = paragraph if not current else f"{current}\n{paragraph}"
            if len(candidate) <= self.max_length:
                current = candidate
                continue
            if current:
                chunks.append(current)
                current = ""
            chunks.extend(self._split_long_line(paragraph))
        if current:
            chunks.append(current)
        return chunks

    def _split_long_line(self, line: str) -> list[str]:
        chunks: list[str] = []
        words = line.split(" ")
        current = ""
        for word in words:
            candidate = word if not current else f"{current} {word}"
            if len(candidate) <= self.max_length:
                current = candidate
                continue
            if current:
                chunks.append(current)
            current = word
            while len(current) > self.max_length:
                chunks.append(current[: self.max_length])
                current = current[self.max_length :]
        if current:
            chunks.append(current)
        return chunks


class ErrorFormatterSkill:
    def unsupported_file(self, extension: str | None) -> str:
        actual = extension or "unknown"
        return (
            f"Формат файла `{actual}` не поддерживается. "
            "Отправьте TXT, MD, DOCX или PDF."
        )

    def file_read_error(self, error: Exception) -> str:
        return f"Не удалось прочитать файл: {error}"

    def clarification_request(self, questions: list[str]) -> str:
        joined = "\n".join(f"{index}. {question}" for index, question in enumerate(questions, 1))
        return f"Нужно уточнить параметры генерации:\n{joined}"

